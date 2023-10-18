
#for code to work here, we need to add last weeks vacacny numbers to the commentary tab, in cell e3
import os
import math
import shutil
import xlwings as xw
import docx as dw
import docx.shared as ds
import datetime as dt

#this input helps us track where we are in the week
on_cycle = input("If you are on-cycle enter 0, if not please enter the number of days missed since Monday ")

#this loop ensures a value is entered
while (on_cycle == ""):
    on_cycle = input("You have not entered a value, please try again ")


#we make this a numebr so we can use it later
on_cycle = int(on_cycle) * -1

today_date = dt.date.today()  + dt.timedelta(on_cycle) - dt.timedelta(2)
file_date = dt.date.strftime(today_date,'%#m-%d-%y')
file_name = "U:/CDNMF/Reports/Unit Availability/2022/Unit Availability " + file_date + ".xlsm"

percent_sequence = ['B', 'C', 'D','F', 'G', 'H','R', 'S', 'T']
ownership_sequence = ['Owner1','Owner2','Owner3']
letters_sequence = ['D','E','H','I','T','U']

UA = xw.Book(file_name, update_links=False)

document = dw.Document()

font = document.styles['Normal'].font
font.name = 'Arial'

document.add_paragraph('Hello Team,')

document.add_paragraph('Please find below the summary of the weekly report prepared by the Analytics team.')

document.add_paragraph(
    'Adjusted vacancy is defined as vacant-unrented rent ready units, also we have mentioned the properties with 100% occupancy and their last date of rental adjustment incase there are any potential rent increases opportunities when the units become vacant.'
)

document.add_paragraph('Please let me know if you have any comments.')


# we have just finished the intro to commentaries, done only once
# next we will begin growth fund1
#
#
#
#first line function writes the first line of the top 5 vacacny
def firstline (para1,current1,old1):
    if (current1>old1):
        para1.add_run('increased relative to last week from ')
        para1.add_run(str(old1))
        para1.add_run('% to ')
        para1.add_run(str(current1))
        para1.add_run('%; top offenders excluding Old Carriage:')
        para1.paragraph_format.left_indent = ds.Inches(0.5)

    elif (current1==old1):
        para1.add_run('remained the same at ')
        para1.add_run(str(current1))
        para1.add_run('%; top offenders excluding Old Carriage:')
        para1.paragraph_format.left_indent = ds.Inches(0.5)


    else:
        para1.add_run('decreased relative to last week from ')
        para1.add_run(str(old1))
        para1.add_run('% to ')
        para1.add_run(str(current1))
        para1.add_run('%; top offenders excluding Old Carriage:')
        para1.paragraph_format.left_indent = ds.Inches(0.5)

#
#
#
# adjtop writes the actual propertty line down and adjusts based on vacancy changes
def adjtop (para,current,old):
    
    if (current>old):
        para.add_run('increased from ')
        para.add_run(str(old))
        para.add_run('% to ')
        para.add_run(str(current))
        para.add_run('%')
        para.paragraph_format.left_indent = ds.Inches(1.0)

    elif (current==old):
        para.add_run('remained the same at ')
        para.add_run(str(current))
        para.add_run('%')
        para.paragraph_format.left_indent = ds.Inches(1.0)


    else:
        para.add_run('decreased from ')
        para.add_run(str(old))
        para.add_run('% to ')
        para.add_run(str(current))
        para.add_run('%')
        para.paragraph_format.left_indent = ds.Inches(1.0)

#
#
#
#
# top5 invokes adjtop, the purpose of this is to alternate column sequences for different ownerships
def top5 (letter1,letter2,letter3):
    
    if ownership == "P":
        i = 5
        while i<11:
            ranger1 = letter1 + str(i)
            ranger2 = letter2 + str(i)
            ranger3 = letter3 + str(i)
            prop1 = str(sheet.range(ranger1).value) + ' '

            g1_cvac1 = sheet.range(ranger2).value
            g1_ovac1 = sheet.range(ranger3).value

            g1_cvac = round(round(g1_cvac1,3)*100,2)
            g1_ovac = round(round(g1_ovac1,3)*100,2)

            p2 = document.add_paragraph(prop1,  style='List Bullet')
            
            # prints the line 
            adjtop(p2,g1_cvac,g1_ovac)
            
            i=i+1
        
    
    else:
        i = 5
        while i<10:
            ranger1 = letter1 + str(i)
            ranger2 = letter2 + str(i)
            ranger3 = letter3 + str(i)
            prop1 = str(sheet.range(ranger1).value) + ' '

            g1_cvac1 = sheet.range(ranger2).value
            g1_ovac1 = sheet.range(ranger3).value

            g1_cvac = round(round(g1_cvac1,3)*100,2)
            g1_ovac = round(round(g1_ovac1,3)*100,2)

            p2 = document.add_paragraph(prop1,  style='List Bullet')
            
            # prints the line 
            adjtop(p2,g1_cvac,g1_ovac)
            
            i=i+1
        

#
#
#
# we are now done functions and begin our first ownership now
#
#
#

# we will use several of our functions below
# the rest of the code will operate in a loop that is lopping through ownerships

#this variable will not change, we only use this sheet for private commentaries
sheet = UA.sheets['Institutional Com.']

w = 0
z = 0
for ownership in ownership_sequence:
    
    p = document.add_paragraph()

    runner = p.add_run(ownership + ":")
    runner.bold = True
    runner.underline = True
    
    ran1 = letters_sequence[w] + '3'
    ran2 = letters_sequence[w+1] + '3'
    
    g1_cvac1 = sheet.range(ran1).value
    g1_ovac1 = sheet.range(ran2).value
    

    g1_cvac = round(round(g1_cvac1,3)*100,2)
    g1_ovac = round(round(g1_ovac1,3)*100,2)

    p1 = document.add_paragraph('Adjusted Vacancy ',  style='List Bullet')
    firstline(p1,g1_cvac,g1_ovac)

    let1 = percent_sequence[z]
    let2 = percent_sequence[z+1]
    let3 = percent_sequence[z+2]
    
    top5(let1, let2, let3)
    
    # we have printed the top 5 vacancy section, we now move to 100% occupancy
    
    occupancy = document.add_paragraph('Properties with 100% occupancy and last date of rental adjustment:',  style='List Bullet')
    occupancy.paragraph_format.left_indent = ds.Inches(0.5)
    
    occ_row = 13
    end = sheet.range(percent_sequence[z]+str(occ_row)).end('down').row
    
    if end == 50:
        end = 13
        
    if end > 51:
        end = 49

    if (end == 13 and sheet.range(percent_sequence[z] + "13").value =="No Properties at 100% Occupancy"):
        
        full_occ = document.add_paragraph("No Properties at 100% Occupancy", style='List Bullet')
        full_occ.paragraph_format.left_indent = ds.Inches(1.0)
        
        end = 12
        
    while occ_row <= end:
        
        occ_range = percent_sequence[z] + str(occ_row)
        
        prop1 = sheet.range(occ_range).value
        #prints out the property name
        all_occ = document.add_paragraph(prop1,  style='List Bullet')
        
        all_occ.add_run(' - ')
        
        
        date_range = percent_sequence[z+1] + str(occ_row)
        date_val = sheet.range(date_range).options(dates = dt.date).value
        
        date_val = dt.datetime.strftime(date_val, "%m/%d/%Y")
        
        all_occ.add_run(str(date_val))
        
        all_occ.paragraph_format.left_indent = ds.Inches(1.0)
        
        occ_row = occ_row + 1
    
        
        # issues: no 1.0 indent present
        # not getting value for end row and not printing
    
    #we now print leasing velocity
    
    velocity = document.add_paragraph("Properties with high leasing velocity (with leases signed in T4W)", style='List Bullet')
    velocity.paragraph_format.left_indent = ds.Inches(0.5)
    
    velocity_row = 51
    velocity_end = sheet.range(percent_sequence[z]+str(velocity_row)).end('down').row
    
    if velocity_end >81:
        velocity_end = 79
    
    if (velocity_end == 80 and sheet.range(percent_sequence[z] + "51").value =="No new leases"):
        
        full_occ = document.add_paragraph("No new units signed", style='List Bullet')
        full_occ.paragraph_format.left_indent = ds.Inches(1.0)
        
        velocity_end = 50
        
        
    while velocity_row <= velocity_end:
        
        velocity_range = percent_sequence[z+1] + str(velocity_row)
        
        prop = sheet.range(percent_sequence[z] + str(velocity_row)).value
        units = int(sheet.range(velocity_range).value)
        #prints out unit number
        all_velocity = document.add_paragraph(str(prop),  style='List Bullet')
        
        all_velocity.add_run(' - ')
        
        all_velocity.add_run(str(units))
        
        all_velocity.add_run(' leases')
        
        all_velocity.paragraph_format.left_indent = ds.Inches(1.0)
        
        velocity_row = velocity_row + 1
    
    
    #finally we print new units rented
    
    unit_number = str(int(sheet.range(percent_sequence[z+2]+"80").value))
    lease = document.add_paragraph(unit_number + ' new units rented this week',  style='List Bullet')
    lease.paragraph_format.left_indent = ds.Inches(0.5)

    lease_row = 81
    lease_end = sheet.range(percent_sequence[z]+str(lease_row)).end('down').row
    
    if lease_end >125:
        lease_end = 81
    
    if (lease_end == 81 and sheet.range(percent_sequence[z] + "81").value =="No new units signed"):
        
        full_occ = document.add_paragraph("No new units signed", style='List Bullet')
        full_occ.paragraph_format.left_indent = ds.Inches(1.0)
        
        lease_end = 80
        
    while lease_row <= lease_end:
        
        lease_range = percent_sequence[z+1] + str(lease_row)
        
        units = int(sheet.range(lease_range).value)
        #prints out unit number
        all_lease = document.add_paragraph(str(units),  style='List Bullet')
        
        all_lease.add_run(' units at ')
        
        prop_range = percent_sequence[z+2] + str(lease_row)
        prop_val = sheet.range(prop_range).value
        
        all_lease.add_run(str(prop_val))
        
        all_lease.paragraph_format.left_indent = ds.Inches(1.0)
        
        lease_row = lease_row + 1
        
    w = w + 2
    z = z + 3

document.save('Institutional Commentaries.docx')

print("see demo file at \"U:\python tester folder\" for your commentaries")

# Prep attachments and parking/high vacancy portion

